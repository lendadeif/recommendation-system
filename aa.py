from docx import Document

# Create the Word document
doc = Document()

# Title
doc.add_heading('RecommendationEngine Class Documentation', level=1)

# Overview
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    'The RecommendationEngine class implements a hybrid recommendation system for a restaurant menu. '
    'It blends association rules, co-purchase similarity, period-based popularity, and freshness signals '
    'to generate up to N recommendations per category (main, side, drink, sauce).'
)

# Dependencies
doc.add_heading('Dependencies', level=2)
dependencies = [
    'pandas',
    'numpy',
    'scipy',
    'scikit-learn',
    'mlxtend',
    'python-docx'
]
for dep in dependencies:
    doc.add_paragraph(f'- {dep}', style='List Bullet')

# Initialization
doc.add_heading('Initialization', level=2)
doc.add_paragraph(
    'Instantiate the engine with historical orders and menu dataframes:\n'
    '```python\n'
    'engine = RecommendationEngine(orders_df, menu_df)\n'
    '```'
)
doc.add_paragraph(
    'The constructor performs:\n'
    '- Date parsing and period assignment (morning/afternoon/evening/night)\n'
    '- Association rule mining (Apriori with support=0.002, confidence>0.15, lift>1.05)\n'
    '- Construction of co-purchase binary vectors\n'
    '- Computation of z-scores for period-based popularity'
)

# Methods
doc.add_heading('Public Methods', level=2)
doc.add_paragraph('recommend(cart_items, top_per_cat=3)')
doc.add_paragraph(
    'Generates recommendations:\n'
    '- Gathers candidate items from rules, period-top sellers, and full menu\n'
    '- Computes component scores: rule_score, time_pop, freshness, content_sim\n'
    '- Combines scores with internal weights (alpha=1.0, beta=0.5, gamma=0.2, sigma=0.3)\n'
    '- Returns a dict mapping each category to a list of recommended item_IDs'
)

# Usage Example
doc.add_heading('Usage Example', level=2)
doc.add_paragraph(
    '```python\n'
    "from recommendation_engine import RecommendationEngine\n\n"
    '# Prepare dataframes\n'
    'orders_df = pd.read_csv("Final.csv")\n'
    'menu_df = pd.read_csv("menu.csv")\n\n'
    'engine = RecommendationEngine(orders_df, menu_df)\n'
    "recs = engine.recommend(['113708', '113709'], top_per_cat=2)\n"
    'print(recs)\n'
    '```'
)

# Configuration & Tuning
doc.add_heading('Configuration & Tuning', level=2)
doc.add_paragraph(
    '- Adjust rule thresholds in the code for different support/confidence/lift\n'
    '- Override internal weights by setting attributes:\n'
    '  engine.alpha = 1.2\n'
    '  engine.beta = 0.6\n'
    '- Modify time periods in _get_period() as needed\n'
)

# Save the document
file_path = '/mnt/data/recommendation_engine.docx'
doc.save(file_path)

file_path
